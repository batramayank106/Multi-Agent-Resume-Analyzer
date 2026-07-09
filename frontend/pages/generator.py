import io
import textwrap
import streamlit as st
from frontend.utils.state import init, page_guide


def _build_tex_report(result: dict) -> str:
    ats = (result or {}).get("ats_result", {}) or {}
    recruiter = (result or {}).get("recruiter_result", {}) or {}
    hm = (result or {}).get("hiring_manager_result", {}) or {}
    sim = (result or {}).get("simulator_results", {}) or {}
    skill_gap = (result or {}).get("skill_gap_result", {}) or {}
    truth = (result or {}).get("truthfulness_result", {}) or {}
    rewrite = (result or {}).get("rewrite_result", {}) or {}
    interview = (result or {}).get("interview_result", {}) or {}
    selection = (result or {}).get("selection_probability", {}) or {}

    tex = r"""\documentclass[11pt]{article}
\usepackage{geometry, hyperref, xcolor, enumitem}
\geometry{margin=1in}
\hypersetup{colorlinks=true,linkcolor=blue,urlcolor=blue}
\definecolor{coral}{HTML}{cc785c}
\definecolor{ink}{HTML}{141413}
\begin{document}
\title{\textbf{\textcolor{coral}{CV Chacha}}\\Multi-Agent Resume Intelligence Report}
\date{\today}
\maketitle
"""
    tex += r"\section*{ATS Score}"
    tex += f"\\textbf{{Overall: {ats.get('overall_score', 'N/A')}/100}}  —  Category: {ats.get('category_label', 'N/A')}\n\n"
    breakdown = ats.get("breakdown", {})
    for cat_key, cat_data in breakdown.items():
        name = cat_key.replace("_", " ").title()
        score = cat_data.get("score", 0)
        tex += f"\\textbf{{{name}}}: {score}/100 \\\\\n"

    if recruiter:
        tex += r"\section*{Recruiter Review}"
        tex += f"Score: {recruiter.get('recruiter_score', 'N/A')}/100  —  Decision: {recruiter.get('decision', 'N/A')}\n\n"
        for dim in recruiter.get("dimensions", []):
            if isinstance(dim, dict):
                tex += f"\\textbf{{{dim.get('name', '')}}}: {dim.get('score', 0)}/100 \\\\\n"
        reasoning = recruiter.get("reasoning", "")
        if reasoning:
            tex += f"\n\\textbf{{Feedback}}: {reasoning[:300]}\\\\\n"

    if hm:
        tex += r"\section*{Hiring Manager Review}"
        tex += f"Score: {hm.get('hiring_manager_score', 'N/A')}/100  —  Decision: {hm.get('decision', 'N/A')}\n\n"
        reasoning = hm.get("reasoning", "")
        if reasoning:
            tex += f"\n\\textbf{{Feedback}}: {reasoning[:300]}\\\\\n"

    if sim:
        tex += r"\section*{Recruiter Simulator}"
        personas = sim.get("personas", sim.get("results", []))
        if isinstance(personas, list):
            for p in personas:
                if isinstance(p, dict):
                    tex += f"\\textbf{{{p.get('persona', p.get('name', 'Persona'))}}}: {p.get('decision', p.get('verdict', 'N/A'))} "
                    tex += f"(Confidence: {p.get('confidence', 0)})\\\\\n"

    if skill_gap:
        tex += r"\section*{Skill Gap Analysis}"
        severity = skill_gap.get("gap_severity")
        if severity:
            tex += f"\\textbf{{Severity}}: {severity}\\\\\n"
        for gap_type in ["missing_required", "missing_preferred"]:
            items = skill_gap.get(gap_type, [])
            if items:
                tex += f"\\textbf{{{gap_type.replace('_', ' ').title()}}}: {', '.join(items)}\\\\\n"
        suggestions = skill_gap.get("suggested_additions", []) or []
        if suggestions:
            tex += f"\\textbf{{Suggested Additions}}: {', '.join(s.get('skill', '') for s in suggestions[:4])}\\\\\n"

    if truth:
        tex += r"\section*{Truthfulness Validation}"
        tex += f"Score: {truth.get('truthfulness_score', 'N/A')}/100 \\\\\n"
        flagged = truth.get("flagged_items", [])
        if flagged:
            tex += f"\\textbf{{Flagged Items}}: {len(flagged)} potential issue(s) found\\\\\n"

    company = (result or {}).get("company_result", {}) or {}
    if company and company.get("company_name"):
        tex += r"\section*{Company Insights}"
        tex += f"\\textbf{{Company}}: {company.get('company_name', 'N/A')}\\\\\n"
        tex += f"\\textbf{{Culture Score}}: {company.get('culture_score', 'N/A')}/100\\\\\n"
        tex += f"\\textbf{{Suitability}}: {company.get('overall_suitability', 'N/A')}\\\\\n"
        pros = company.get("pros", [])
        if pros:
            tex += f"\\textbf{{Pros}}: {', '.join(pros[:3])}\\\\\n"
        cons = company.get("cons", [])
        if cons:
            tex += f"\\textbf{{Cons}}: {', '.join(cons[:3])}\\\\\n"

    salary = (result or {}).get("salary_result", {}) or {}
    if salary and salary.get("salary_range"):
        tex += r"\section*{Salary Estimation}"
        tex += f"\\textbf{{Range}}: {salary.get('salary_range', 'N/A')}\\\\\n"
        tex += f"\\textbf{{Confidence}}: {salary.get('confidence', 'N/A')}\\\\\n"
        factors = salary.get("factors", [])
        if factors:
            tex += f"\\textbf{{Factors}}: {', '.join(factors[:4])}\\\\\n"

    if rewrite:
        tex += r"\section*{Resume Rewrite Suggestions}"
        rewritten = rewrite.get("rewritten_resume", "")
        if rewritten and "API key" not in rewritten:
            tex += "\\begin{verbatim}\n" + rewritten[:2000] + "\n\\end{verbatim}\n"

    if interview:
        tex += r"\section*{Interview Questions}"
        questions = interview.get("questions", [])
        if isinstance(questions, list):
            for q in questions[:5]:
                if isinstance(q, dict):
                    tex += f"\\textbf{{[{q.get('difficulty', 'Medium')}]}} {q.get('question', '')}\\\\\n"

    if selection:
        tex += r"\section*{Selection Probability}"
        overall = selection.get("overall_probability")
        if overall is not None:
            tex += f"\\textbf{{Overall Probability: {overall}\\%}}\\\\\n\n"
        stages = selection.get("stage_probabilities", {})
        if isinstance(stages, dict):
            for stage, prob in stages.items():
                prob_str = f"{prob}\%" if isinstance(prob, (int, float)) else str(prob)
                tex += f"\\textbf{{{stage}}}: {prob_str}\\\\\n"
        key_factors = selection.get("key_factors", [])
        if key_factors:
            tex += "\n\\textbf{Key Factors:}\\\\\n"
            for f in key_factors[:4]:
                tex += f"  • {f}\\\\\n"
        recommendations = selection.get("recommendations", [])
        if recommendations:
            tex += "\n\\textbf{Recommendations:}\\\\\n"
            for r in recommendations[:3]:
                tex += f"  • {r}\\\\\n"

    tex += r"\end{document}"
    return tex





def _build_executive_summary(result: dict) -> str:
    ats = (result or {}).get("ats_result", {}) or {}
    recruiter = (result or {}).get("recruiter_result", {}) or {}
    hm = (result or {}).get("hiring_manager_result", {}) or {}
    truth = (result or {}).get("truthfulness_result", {}) or {}
    selection = (result or {}).get("selection_probability", {}) or {}
    skill_gap = (result or {}).get("skill_gap_result", {}) or {}

    lines = []
    lines.append("=" * 60)
    lines.append("  CV CHACHA - EXECUTIVE SUMMARY")
    lines.append("  Multi-Agent Resume Intelligence Report")
    lines.append("=" * 60)
    lines.append("")

    lines.append(f"  ATS Score:             {ats.get('overall_score', 'N/A')}/100 ({ats.get('category_label', 'N/A')})")
    lines.append(f"  Recruiter Score:       {recruiter.get('recruiter_score', 'N/A')}/100 ({recruiter.get('decision', 'N/A')})")
    rec_reasoning = recruiter.get("reasoning", "")
    if rec_reasoning:
        lines.append(f"  Recruiter Feedback:    {rec_reasoning[:200]}")
    lines.append(f"  Hiring Manager Score:  {hm.get('hiring_manager_score', 'N/A')}/100 ({hm.get('decision', 'N/A')})")
    hm_reasoning = hm.get("reasoning", "")
    if hm_reasoning:
        lines.append(f"  HM Feedback:           {hm_reasoning[:200]}")
    lines.append(f"  Truthfulness Score:    {truth.get('truthfulness_score', 'N/A')}/100")
    flagged = truth.get("flagged_items", [])
    if flagged:
        lines.append(f"  Flagged Items:         {len(flagged)} potential issue(s) found")
    lines.append("")

    overall = selection.get("overall_probability")
    if overall is not None:
        lines.append(f"  Overall Selection Probability: {overall}%")
    selection_stages = selection.get("stage_probabilities", {})
    if isinstance(selection_stages, dict) and selection_stages:
        lines.append("  Stage Probabilities:")
        for stage, prob in selection_stages.items():
            prob_str = f"{prob}%" if isinstance(prob, (int, float)) else str(prob)
            lines.append(f"    {stage}: {prob_str}")
    key_factors = selection.get("key_factors", [])
    if key_factors:
        lines.append("  Key Factors:")
        for f in key_factors[:3]:
            lines.append(f"    • {f}")
    recommendations = selection.get("recommendations", [])
    if recommendations:
        lines.append("  Recommendations:")
        for r in recommendations[:3]:
            lines.append(f"    • {r}")

    lines.append("")
    missing_skills = skill_gap.get("missing_required", []) or []
    if missing_skills:
        lines.append(f"  Missing Required Skills: {', '.join(missing_skills[:8])}")
    severity = skill_gap.get("gap_severity")
    if severity:
        lines.append(f"  Gap Severity: {severity}")
    suggestions = skill_gap.get("suggested_additions", []) or []
    if suggestions:
        lines.append(f"  Suggested Additions: {', '.join(s.get('skill', '') for s in suggestions[:4])}")

    lines.append("")
    lines.append("-" * 60)
    lines.append("  Generated by CV Chacha - https://github.com/batramayank106/Multi-Agent-Resume-Analyzer")
    lines.append("=" * 60)

    return "\n".join(lines)


def _build_docx_bytes(result: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    doc.add_heading("CV Chacha", level=1)
    doc.add_heading("Multi-Agent Resume Intelligence Report", level=2)
    doc.add_paragraph("")

    ats = (result or {}).get("ats_result", {}) or {}
    recruiter = (result or {}).get("recruiter_result", {}) or {}
    hm = (result or {}).get("hiring_manager_result", {}) or {}
    sim = (result or {}).get("simulator_results", {}) or {}
    skill_gap = (result or {}).get("skill_gap_result", {}) or {}
    truth = (result or {}).get("truthfulness_result", {}) or {}
    rewrite = (result or {}).get("rewrite_result", {}) or {}
    interview = (result or {}).get("interview_result", {}) or {}
    selection = (result or {}).get("selection_probability", {}) or {}

    doc.add_heading("ATS Score", level=2)
    doc.add_paragraph(f"Overall: {ats.get('overall_score', 'N/A')}/100  —  Category: {ats.get('category_label', 'N/A')}")
    for cat_key, cat_data in (ats.get("breakdown", {}) or {}).items():
        name = cat_key.replace("_", " ").title()
        score = cat_data.get("score", 0)
        doc.add_paragraph(f"  {name}: {score}/100", style="List Bullet")

    if recruiter:
        doc.add_heading("Recruiter Review", level=2)
        doc.add_paragraph(f"Score: {recruiter.get('recruiter_score', 'N/A')}/100  —  Decision: {recruiter.get('decision', 'N/A')}")
        reasoning = recruiter.get("reasoning", "")
        if reasoning:
            doc.add_paragraph(f"Feedback: {reasoning[:300]}")

    if hm:
        doc.add_heading("Hiring Manager Review", level=2)
        doc.add_paragraph(f"Score: {hm.get('hiring_manager_score', 'N/A')}/100  —  Decision: {hm.get('decision', 'N/A')}")
        reasoning = hm.get("reasoning", "")
        if reasoning:
            doc.add_paragraph(f"Feedback: {reasoning[:300]}")

    if sim:
        doc.add_heading("Recruiter Simulator", level=2)
        personas = sim.get("personas", sim.get("results", []))
        for p in (personas or []):
            if isinstance(p, dict):
                name = p.get("persona", p.get("name", "Persona"))
                decision = p.get("decision", p.get("verdict", "N/A"))
                conf = p.get("confidence", 0)
                doc.add_paragraph(f"{name}: {decision} (confidence: {conf})", style="List Bullet")

    if skill_gap:
        doc.add_heading("Skill Gap Analysis", level=2)
        severity = skill_gap.get("gap_severity")
        if severity:
            p = doc.add_paragraph()
            run = p.add_run("Severity: ")
            run.bold = True
            p.add_run(severity)
        for gap_type in ["missing_required", "missing_preferred"]:
            items = skill_gap.get(gap_type, []) or []
            if items:
                p = doc.add_paragraph()
                run = p.add_run(f"{gap_type.replace('_', ' ').title()}: ")
                run.bold = True
                p.add_run(", ".join(items[:10]))
        suggestions = skill_gap.get("suggested_additions", []) or []
        if suggestions:
            p = doc.add_paragraph()
            run = p.add_run("Suggested Additions: ")
            run.bold = True
            p.add_run(", ".join(s.get("skill", "") for s in suggestions[:4]))

    if truth:
        doc.add_heading("Truthfulness Validation", level=2)
        doc.add_paragraph(f"Score: {truth.get('truthfulness_score', 'N/A')}/100")
        flagged = truth.get("flagged_items", [])
        if flagged:
            doc.add_paragraph(f"Flagged Items: {len(flagged)} potential issue(s) found")

    company_doc = (result or {}).get("company_result", {}) or {}
    if company_doc and company_doc.get("company_name"):
        doc.add_heading("Company Insights", level=2)
        doc.add_paragraph(f"Company: {company_doc.get('company_name', 'N/A')}")
        doc.add_paragraph(f"Culture Score: {company_doc.get('culture_score', 'N/A')}/100")
        doc.add_paragraph(f"Suitability: {company_doc.get('overall_suitability', 'N/A')}")
        pros = company_doc.get("pros", [])
        if pros:
            p = doc.add_paragraph()
            run = p.add_run("Pros: ")
            run.bold = True
            p.add_run("; ".join(pros[:3]))
        cons = company_doc.get("cons", [])
        if cons:
            p = doc.add_paragraph()
            run = p.add_run("Cons: ")
            run.bold = True
            p.add_run("; ".join(cons[:3]))

    salary_doc = (result or {}).get("salary_result", {}) or {}
    if salary_doc and salary_doc.get("salary_range"):
        doc.add_heading("Salary Estimation", level=2)
        doc.add_paragraph(f"Range: {salary_doc.get('salary_range', 'N/A')} ({salary_doc.get('currency', 'N/A')})")
        doc.add_paragraph(f"Confidence: {salary_doc.get('confidence', 'N/A')}")
        factors = salary_doc.get("factors", [])
        if factors:
            p = doc.add_paragraph()
            run = p.add_run("Factors: ")
            run.bold = True
            p.add_run("; ".join(factors[:4]))

    if rewrite:
        doc.add_heading("Resume Rewrite", level=2)
        rewritten = rewrite.get("rewritten_resume", "")
        if rewritten and "API key" not in rewritten:
            doc.add_paragraph(rewritten[:2000])

    if interview:
        doc.add_heading("Interview Questions", level=2)
        questions = interview.get("questions", []) or []
        for q in questions[:8]:
            if isinstance(q, dict):
                diff = q.get("difficulty", "Medium")
                question = q.get("question", "")
                doc.add_paragraph(f"[{diff}] {question}", style="List Bullet")

    if selection:
        doc.add_heading("Selection Probability", level=2)
        overall = selection.get("overall_probability")
        if overall is not None:
            doc.add_paragraph(f"Overall Probability: {overall}%")
        stages = selection.get("stage_probabilities", {})
        if isinstance(stages, dict):
            for stage, prob in stages.items():
                prob_str = f"{prob}%" if isinstance(prob, (int, float)) else str(prob)
                doc.add_paragraph(f"{stage}: {prob_str}", style="List Bullet")
        key_factors = selection.get("key_factors", [])
        if key_factors:
            p = doc.add_paragraph()
            run = p.add_run("Key Factors: ")
            run.bold = True
            p.add_run("; ".join(key_factors[:4]))
        recommendations = selection.get("recommendations", [])
        if recommendations:
            p = doc.add_paragraph()
            run = p.add_run("Recommendations: ")
            run.bold = True
            p.add_run("; ".join(recommendations[:3]))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render():
    init()
    st.markdown(
        "<h1 style='font-size: 2.2rem;'>Resume Generator</h1>",
        unsafe_allow_html=True,
    )
    page_guide(
        "Resume Generator & Reports",
        "Download your analysis results in multiple formats: optimized resume (TXT), executive summary (TXT), LaTeX report (TEX), Word report (DOCX), and full report (TXT).",
        "Buttons to download each format appear after analysis. The optimized resume is a rewritten version from the Rewrite agent. Executive summary condenses all scores into one page. LaTeX/Word/TXT contain the full multi-agent report.",
        "Run a full analysis first. Click any download button to save. The rewritten resume is an LLM draft — always review and customize before using in applications.",
    )
    result = st.session_state.get("analysis_result")
    rewrite = (result or {}).get("rewrite_result", {}) if result else None

    if not rewrite:
        st.info("Run a full analysis first to generate output files.")
        return

    rewritten = rewrite.get("rewritten_resume", "")
    has_content = bool(rewritten and "API key" not in rewritten)

    if not has_content:
        st.warning("Set up HF_API_KEY in your .env file to generate optimized resumes.")
        return

    exec_summary = _build_executive_summary(result)

    st.markdown("<hr style='margin: 0.5rem 0;'>", unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.download_button(
            label="Download Optimized Resume (TXT)",
            data=rewritten,
            file_name="optimized_resume.txt",
            mime="text/plain",
            type="primary",
            use_container_width=True,
        )
    with col2:
        st.download_button(
            label="Download Executive Summary (TXT)",
            data=exec_summary,
            file_name="cv_chacha_executive_summary.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with col3:
        tex_content = _build_tex_report(result)
        st.download_button(
            label="Download LaTeX Report (TEX)",
            data=tex_content,
            file_name="cv_chacha_report.tex",
            mime="text/x-tex",
            use_container_width=True,
        )
    with col4:
        try:
            docx_bytes = _build_docx_bytes(result)
            st.download_button(
                label="Download Word Report (DOCX)",
                data=docx_bytes,
                file_name="cv_chacha_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"DOCX generation failed: {e}")

    st.markdown("<br>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        full_report = exec_summary + "\n\n" + "=" * 60 + "\nOPTIMIZED RESUME\n" + "=" * 60 + "\n\n" + rewritten
        st.download_button(
            label="Download Full Report (TXT)",
            data=full_report,
            file_name="cv_chacha_full_report.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with col3:
        st.markdown("<div style='padding-top: 0.4rem;'></div>", unsafe_allow_html=True)
        st.markdown(
            "<p style='color: #6c6a64; font-size: 0.8rem; text-align: center;'>"
            "DOCX · TEX · TXT</p>",
            unsafe_allow_html=True,
        )

    st.markdown("<hr style='margin: 1.5rem 0;'>", unsafe_allow_html=True)

    tabs = st.tabs(["Optimized Resume", "Executive Summary", "Full Report"])
    with tabs[0]:
        st.markdown(
            f"<div style='background: #f5f0e8; border: 1px solid #e6dfd8; border-radius: 8px; padding: 1rem; "
            f"font-family: JetBrains Mono, monospace; font-size: 0.8rem; color: #3d3d3a; "
            f"white-space: pre-wrap; max-height: 500px; overflow-y: auto;'>{rewritten}</div>",
            unsafe_allow_html=True,
        )
    with tabs[1]:
        st.markdown(
            f"<div style='background: #f5f0e8; border: 1px solid #e6dfd8; border-radius: 8px; padding: 1rem; "
            f"font-family: JetBrains Mono, monospace; font-size: 0.8rem; color: #3d3d3a; "
            f"white-space: pre-wrap; max-height: 400px; overflow-y: auto;'>{exec_summary}</div>",
            unsafe_allow_html=True,
        )
    with tabs[2]:
        st.markdown(
            f"<div style='background: #f5f0e8; border: 1px solid #e6dfd8; border-radius: 8px; padding: 1rem; "
            f"font-family: JetBrains Mono, monospace; font-size: 0.8rem; color: #3d3d3a; "
            f"white-space: pre-wrap; max-height: 500px; overflow-y: auto;'>{full_report}</div>",
            unsafe_allow_html=True,
        )
